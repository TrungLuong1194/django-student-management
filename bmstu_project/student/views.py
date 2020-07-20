from django.shortcuts import render, redirect
from django.http import HttpResponse
from student.models import Major, UserProfile
from student.forms import MajorForm, UserProfileForm
from django.contrib.auth.decorators import login_required
from registration.backends.simple.views import RegistrationView
from django.urls import reverse
from django.contrib.auth.models import User
from docx import *
from io import BytesIO
from datetime import date
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Inches


def index(request):
	context_dict = {'boldmessage': "Crunchy, creamy, cookie, candy, cupcake!"}

	return render(request, 'student/index.html', context=context_dict)

@login_required
def majors_list(request):
	majors_list = Major.objects.all()

	return render(request, 'student/majors_list.html',
		{'majors_list' : majors_list})

@login_required
def add_major(request):
    form = MajorForm()

    if request.method == 'POST':
        form = MajorForm(request.POST)

        if form.is_valid():
            form.save(commit=True)

            return index(request)
        else:
            print(form.errors)

    return render(request, 'student/add_major.html', {'form': form})

@login_required
def register_profile(request):
    form = UserProfileForm()

    if request.method == 'POST':
        form = UserProfileForm(request.POST)
        if form.is_valid():
            user_profile = form.save(commit=False)
            user_profile.user = request.user
            user_profile.save()

            return redirect('index')
        else:
            print(form.errors)

    context_dict = {'form':form}

    return render(request, 'student/profile_registration.html', context_dict)


class StudentRegistrationView(RegistrationView):
    def get_success_url(self, user):
        return reverse('register_profile')

@login_required
def profile(request, username):
    try:
        user = User.objects.get(username=username)
    except User.DoesNotExist:
        return redirect('index')
    
    userprofile = UserProfile.objects.get_or_create(user=user)[0]
    form = UserProfileForm({
            'major': userprofile.major,
            'fullName': userprofile.fullName,
            'gender': userprofile.gender,
            'birthday': userprofile.birthday,
            'ethnic': userprofile.ethnic,
            'religion': userprofile.religion,
            'studyYear': userprofile.studyYear,
            'addressVN': userprofile.addressVN,
            'addressRu': userprofile.addressRu,
            'phone': userprofile.phone,
            'workPlace': userprofile.workPlace,
            'dateOfAdmission': userprofile.dateOfAdmission,
            'dateOfStudy': userprofile.dateOfStudy,
            'timeOfStudy': userprofile.timeOfStudy,
            'infoOfStudy': userprofile.infoOfStudy,

            'ruSubject1': userprofile.ruSubject1,
            'viSubject1': userprofile.viSubject1,
            'resultSubject1': userprofile.resultSubject1,

            'ruSubject2': userprofile.ruSubject2,
            'viSubject2': userprofile.viSubject2,
            'resultSubject2': userprofile.resultSubject2,

            'ruSubject3': userprofile.ruSubject3,
            'viSubject3': userprofile.viSubject3,
            'resultSubject3': userprofile.resultSubject3,

            'ruSubject4': userprofile.ruSubject4,
            'viSubject4': userprofile.viSubject4,
            'resultSubject4': userprofile.resultSubject4,

            'ruSubject5': userprofile.ruSubject5,
            'viSubject5': userprofile.viSubject5,
            'resultSubject5': userprofile.resultSubject5,

            'ruSubject6': userprofile.ruSubject6,
            'viSubject6': userprofile.viSubject6,
            'resultSubject6': userprofile.resultSubject6,

            'ruSubject7': userprofile.ruSubject7,
            'viSubject7': userprofile.viSubject7,
            'resultSubject7': userprofile.resultSubject7,

            'ruSubject8': userprofile.ruSubject8,
            'viSubject8': userprofile.viSubject8,
            'resultSubject8': userprofile.resultSubject8,

            'ruSubject9': userprofile.ruSubject9,
            'viSubject9': userprofile.viSubject9,
            'resultSubject9': userprofile.resultSubject9,

            'ruSubject10': userprofile.ruSubject10,
            'viSubject10': userprofile.viSubject10,
            'resultSubject10': userprofile.resultSubject10,

            'ruSubject11': userprofile.ruSubject11,
            'viSubject11': userprofile.viSubject11,
            'resultSubject11': userprofile.resultSubject11,

            'ruSubject12': userprofile.ruSubject12,
            'viSubject12': userprofile.viSubject12,
            'resultSubject12': userprofile.resultSubject12,

            'nameBank': userprofile.nameBank,
            'nameAccount': userprofile.nameAccount,
        })
    
    if request.method == 'POST':
        form = UserProfileForm(request.POST, instance=userprofile)
        if form.is_valid():
            form.save(commit=True)
            return redirect('profile', user.username)
        else:
            print(form.errors)
    
    return render(request, 'student/profile.html', 
            {'userprofile': userprofile, 'selecteduser': user, 'form': form})

@login_required
def profiles_list(request):
    userprofile_list = UserProfile.objects.all()

    return render(request, 'student/profiles_list.html',
            {'userprofile_list' : userprofile_list})

@login_required
def report_document(request):
    user_current = None
    if request.user.is_authenticated:
        user_current = UserProfile.objects.get(user=request.user.id)

    document = Document()
    docx_title='Bao cao hoc tap.docx'

    # Setting document
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Content
    paragraph1 = document.add_paragraph()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM').bold = True
    paragraph1.add_run('\nĐộc lập - Tự do - Hạnh phúc').bold = True
    paragraph1.add_run('\n\nBÁO CÁO TIẾN ĐỘ HỌC TẬP').bold = True
    paragraph1.add_run('\n(từ ngày  tháng  năm 20   đến ngày  tháng  năm 20  )')
    paragraph1.add_run('\n\nKính gửi: Bộ Giáo dục và Đào tạo')

    paragraph2 = document.add_paragraph()
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph2.add_run('1. Họ và tên: ')
    paragraph2.add_run(user_current.fullName).bold = True
    paragraph2.add_run('\t\t\t\t\tNam/nữ: ')
    paragraph2.add_run(user_current.gender).bold = True

    paragraph2.add_run('\n2. Ngày sinh: ')
    paragraph2.add_run(user_current.birthday.strftime('%d/%m/%Y')).bold = True

    paragraph2.add_run('\n3. Dân tộc: ')
    paragraph2.add_run(user_current.ethnic).bold = True
    paragraph2.add_run('\t\t\t\t\t\t\tTôn giáo: ')
    paragraph2.add_run(user_current.religion).bold = True

    paragraph2.add_run('\n4. Năm trúng tuyển (đối với lưu học sinh học bổng): ')
    paragraph2.add_run(str(user_current.studyYear)).bold = True
    paragraph2.add_run('. Năm đi học: ')
    paragraph2.add_run(str(user_current.studyYear)).bold = True

    paragraph2.add_run('\n5. Địa chỉ nơi ở tại Việt Nam: ')
    paragraph2.add_run(user_current.addressVN).bold = True

    paragraph2.add_run('\n6. Cơ quan công tác (nếu có): ')
    if user_current.workPlace:
        paragraph2.add_run(user_current.workPlace).bold = True

    paragraph2.add_run('\n7. Diện học bổng (Hiệp định/NSNN/Khác, ghi cụ thể): ')
    paragraph2.add_run('Hiệp định').bold = True

    paragraph2.add_run('\n8. Ngành học ở nước ngoài (ghi tiếng Việt và tiếng Anh):\n')
    paragraph2.add_run('    ' + user_current.major.viName).bold = True
    paragraph2.add_run(' (' + user_current.major.enName + ')').bold = True

    paragraph2.add_run('\n9. Tên và địa chỉ trường học ở nước ngoài (ghi tiếng Việt và tiếng Anh): ')
    paragraph2.add_run('Đại học Kỹ thuật Quốc gia Mát-xcơ-va mang tên N.E. Bauman - Bauman Moscow State Technical University\n\
Địa chỉ: số 5, đường Baumanskaya-2, Mát-xcơ-va, 105005').bold = True

    paragraph2.add_run('\n10. Ngày đến trường nhập học: ')
    paragraph2.add_run(user_current.dateOfAdmission.strftime('%d/%m/%Y')).bold = True

    paragraph2.add_run('\n11. Ngày bắt đầu khóa học (theo thông báo của trường): ')
    paragraph2.add_run(user_current.dateOfStudy.strftime('%d/%m/%Y')).bold = True

    paragraph2.add_run('\n12. Thời gian đào tạo (theo thông báo của trường): ')
    paragraph2.add_run(user_current.timeOfStudy).bold = True

    paragraph2.add_run('\n13. Đang học học kỳ mấy, thời gian còn lại: ')
    paragraph2.add_run(user_current.infoOfStudy).bold = True

    paragraph2.add_run('\n14. Địa chỉ nơi ở nước ngoài: ')
    paragraph2.add_run(user_current.addressRu).bold = True

    paragraph2.add_run('\n15. E-mail ở nước ngoài: ')
    paragraph2.add_run(request.user.email).bold = True

    paragraph2.add_run('\n16. Điện thoại liên hệ ở nước ngoài: ')
    paragraph2.add_run(str(user_current.phone)).bold = True

    GRADE_CHOICES = {
        "Отлично": "Giỏi",
        "Хорошо": "Khá",
        "Удовлетворительно": "Trung Bình",
        "Зачтено": "Đạt"
    }

    paragraph2.add_run('\n17. Kết quả học tập:\n')
    if user_current.ruSubject1:
        paragraph2.add_run('- ' + user_current.ruSubject1).bold = True
        paragraph2.add_run(' (' + user_current.viSubject1 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject1 + ' (' + GRADE_CHOICES[user_current.resultSubject1] + ')')

    if user_current.ruSubject2:
        paragraph2.add_run('\n- ' + user_current.ruSubject2).bold = True
        paragraph2.add_run(' (' + user_current.viSubject2 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject2 + ' (' + GRADE_CHOICES[user_current.resultSubject2] + ')')

    if user_current.ruSubject3:
        paragraph2.add_run('\n- ' + user_current.ruSubject3).bold = True
        paragraph2.add_run(' (' + user_current.viSubject3 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject3 + ' (' + GRADE_CHOICES[user_current.resultSubject3] + ')')

    if user_current.ruSubject4:
        paragraph2.add_run('\n- ' + user_current.ruSubject4).bold = True
        paragraph2.add_run(' (' + user_current.viSubject4 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject4 + ' (' + GRADE_CHOICES[user_current.resultSubject4] + ')')

    if user_current.ruSubject5:
        paragraph2.add_run('\n- ' + user_current.ruSubject5).bold = True
        paragraph2.add_run(' (' + user_current.viSubject5 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject5 + ' (' + GRADE_CHOICES[user_current.resultSubject5] + ')')

    if user_current.ruSubject6:
        paragraph2.add_run('\n- ' + user_current.ruSubject6).bold = True
        paragraph2.add_run(' (' + user_current.viSubject6 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject6 + ' (' + GRADE_CHOICES[user_current.resultSubject6] + ')')

    if user_current.ruSubject7:
        paragraph2.add_run('\n- ' + user_current.ruSubject7).bold = True
        paragraph2.add_run(' (' + user_current.viSubject7 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject7 + ' (' + GRADE_CHOICES[user_current.resultSubject7] + ')')

    if user_current.ruSubject8:
        paragraph2.add_run('\n- ' + user_current.ruSubject8).bold = True
        paragraph2.add_run(' (' + user_current.viSubject8 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject8 + ' (' + GRADE_CHOICES[user_current.resultSubject8] + ')')

    if user_current.ruSubject9:
        paragraph2.add_run('\n- ' + user_current.ruSubject9).bold = True
        paragraph2.add_run(' (' + user_current.viSubject9 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject9 + ' (' + GRADE_CHOICES[user_current.resultSubject9] + ')')

    if user_current.ruSubject10:
        paragraph2.add_run('\n- ' + user_current.ruSubject10).bold = True
        paragraph2.add_run(' (' + user_current.viSubject10 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject10 + ' (' + GRADE_CHOICES[user_current.resultSubject10] + ')')

    if user_current.ruSubject11:
        paragraph2.add_run('\n- ' + user_current.ruSubject11).bold = True
        paragraph2.add_run(' (' + user_current.viSubject11 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject11 + ' (' + GRADE_CHOICES[user_current.resultSubject11] + ')')

    if user_current.ruSubject12:
        paragraph2.add_run('\n- ' + user_current.ruSubject12).bold = True
        paragraph2.add_run(' (' + user_current.viSubject12 + '):').bold = True
        paragraph2.add_run('\n  ' + user_current.resultSubject12 + ' (' + GRADE_CHOICES[user_current.resultSubject12] + ')')

    paragraph2.add_run('\n18. Họ tên người hướng dẫn (supervisor) hoặc người tư vấn (adviser):.....')
    paragraph2.add_run('\nĐịa chỉ e-mail của người hướng dẫn/tư vấn:.....')

    paragraph2.add_run('\n19. Kiến nghị, đề xuất (nếu có):.....')

    paragraph2.add_run('\n20. Đề nghị cấp học phí, sinh hoạt phí (đối với lưu học sinh học bổng):')
    paragraph2.add_run('\nĐã nhận sinh hoạt phí đến hết tháng  năm 20  ')
    paragraph2.add_run('\nHọc kỳ cuối cùng xin được chuyển sinh hoạt phí đến hết tháng  năm 20  , tổng số 6 tháng.')
    paragraph2.add_run('\nCập nhật số tài khoản cá nhân đã đăng ký:')

    table1 = document.add_table(rows=1, cols=3, style='Table Grid')
    row = table1.rows[0]

    row.cells[0].paragraphs[0].add_run('- Tên ngân hàng:\n')
    row.cells[0].paragraphs[0].add_run(user_current.nameBank).bold = True
    row.cells[0].paragraphs[0].add_run('\n- Địa chỉ ngân hàng:')
    row.cells[0].paragraphs[0].add_run('\nMoscow, Russia').bold = True
    row.cells[0].paragraphs[0].add_run('\n- Mã số ngân hàng (Swift Code):\n\n')
    row.cells[0].paragraphs[0].add_run('- Thông tin ngân hàng trung gian (nếu có):\n')

    row.cells[1].paragraphs[0].add_run('- Tên người hưởng:\n')
    row.cells[1].paragraphs[0].add_run('(chủ tài khoản cá nhân):\n')
    row.cells[1].paragraphs[0].add_run(user_current.nameAccount).bold = True
    row.cells[1].paragraphs[0].add_run('\n- Địa chỉ người hưởng:')
    row.cells[1].paragraphs[0].add_run('\nMoscow, Russia').bold = True

    row.cells[2].paragraphs[0].add_run('- Số tài khoản:\n\n')
    row.cells[2].paragraphs[0].add_run('- Số Iban (nếu có):')

    document.add_paragraph()

    table2 = document.add_table(rows=1, cols=2)
    row = table2.rows[0]

    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].paragraphs[0].add_run('Xác nhận của đơn vị\nĐơn vị trưởng').bold = True

    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[1].paragraphs[0].add_run('Mát-xcơ-va, ngày  tháng  năm 20  ')
    row.cells[1].paragraphs[0].add_run('\nNgười báo cáo').bold = True
    row.cells[1].paragraphs[0].add_run('\n\n\n\n\n' + user_current.fullName).bold = True

    # Prepare document for download        
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length

    return response

@login_required
def grade_document(request):
    user_current = None
    if request.user.is_authenticated:
        user_current = UserProfile.objects.get(user=request.user.id)

    document = Document()
    docx_title="Ban dich bang diem.docx"

    # Setting document
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Content
    table1 = document.add_table(rows=1, cols=2)
    table1.columns[0].width=Inches(4)
    table1.columns[1].width=Inches(2)
    row = table1.rows[0]

    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].paragraphs[0].add_run('BỘ GIÁO DỤC VÀ KHOA HỌC')
    row.cells[0].paragraphs[0].add_run('\nLIÊN BANG NGA')
    row.cells[0].paragraphs[0].add_run('\nCục Chính sách Nhà nước trong Giáo dục Đại học')
    row.cells[0].paragraphs[0].add_run('\nCơ sở giáo dục đại học ngân sách nhà nước Liên bang')
    row.cells[0].paragraphs[0].add_run('\n« Đại học Kỹ thuật Quốc gia Moskva mang tên N.E.Bauman')
    row.cells[0].paragraphs[0].add_run('\n(Đại học Nghiên cứu Quốc gia) »')
    row.cells[0].paragraphs[0].add_run('\nHệ chương trình đào tạo quốc tế')
    row.cells[0].paragraphs[0].add_run('\n105005, Nhà số 5, đường số 2 Baumanskaya, Moskva')
    row.cells[0].paragraphs[0].add_run('\nSđt: (499) 267-00-82, Fax: 267-48-44. OKPO')
    row.cells[0].paragraphs[0].add_run('\n02066434B OGRN 1027739051779 YNN/KPP')
    row.cells[0].paragraphs[0].add_run('\n7701002520/770101001')
    row.cells[0].paragraphs[0].add_run('\n... số ...')

    paragraph1 = document.add_paragraph()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph1.add_run('\nCHỨNG NHẬN').bold = True
    paragraph1.add_run('\nVề kết quả học tập').bold = True

    paragraph2 = document.add_paragraph()
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph2.add_run('\n\tCấp cho sinh viên lớp ... ' + user_current.fullName + ', ngày sinh ' + user_current.birthday.strftime('%d.%m.%Y') +\
        ', đang học năm ... tại trường THKT mang tên Bauman về kết quả học tập: Không có môn học bị nợ, tham dự tất các các tiết học đều đặn.')

    paragraph3 = document.add_paragraph()
    paragraph3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph3.add_run('Kết quả học tập kì học thứ ...:\n').bold = True

    GRADE_CHOICES = {
        "Отлично": "Giỏi",
        "Хорошо": "Khá",
        "Удовлетворительно": "Trung Bình",
        "Зачтено": "Đạt"
    }
    
    if user_current.ruSubject1:
        paragraph3.add_run('- ' + user_current.ruSubject1)
        paragraph3.add_run(' (' + user_current.viSubject1 + '): ')
        paragraph3.add_run(user_current.resultSubject1 + ' (' + GRADE_CHOICES[user_current.resultSubject1] + ')').bold = True

    if user_current.ruSubject2:
        paragraph3.add_run('\n- ' + user_current.ruSubject2)
        paragraph3.add_run(' (' + user_current.viSubject2 + '): ')
        paragraph3.add_run(user_current.resultSubject2 + ' (' + GRADE_CHOICES[user_current.resultSubject2] + ')').bold = True

    if user_current.ruSubject3:
        paragraph3.add_run('\n- ' + user_current.ruSubject3)
        paragraph3.add_run(' (' + user_current.viSubject3 + '): ')
        paragraph3.add_run(user_current.resultSubject3 + ' (' + GRADE_CHOICES[user_current.resultSubject3] + ')').bold = True

    if user_current.ruSubject4:
        paragraph3.add_run('\n- ' + user_current.ruSubject4)
        paragraph3.add_run(' (' + user_current.viSubject4 + '): ')
        paragraph3.add_run(user_current.resultSubject4 + ' (' + GRADE_CHOICES[user_current.resultSubject4] + ')').bold = True

    if user_current.ruSubject5:
        paragraph3.add_run('\n- ' + user_current.ruSubject5)
        paragraph3.add_run(' (' + user_current.viSubject5 + '): ')
        paragraph3.add_run(user_current.resultSubject5 + ' (' + GRADE_CHOICES[user_current.resultSubject5] + ')').bold = True

    if user_current.ruSubject6:
        paragraph3.add_run('\n- ' + user_current.ruSubject6)
        paragraph3.add_run(' (' + user_current.viSubject6 + '): ')
        paragraph3.add_run(user_current.resultSubject6 + ' (' + GRADE_CHOICES[user_current.resultSubject6] + ')').bold = True

    if user_current.ruSubject7:
        paragraph3.add_run('\n- ' + user_current.ruSubject7)
        paragraph3.add_run(' (' + user_current.viSubject7 + '): ')
        paragraph3.add_run(user_current.resultSubject7 + ' (' + GRADE_CHOICES[user_current.resultSubject7] + ')').bold = True

    if user_current.ruSubject8:
        paragraph3.add_run('\n- ' + user_current.ruSubject8)
        paragraph3.add_run(' (' + user_current.viSubject8 + '): ')
        paragraph3.add_run(user_current.resultSubject8 + ' (' + GRADE_CHOICES[user_current.resultSubject8] + ')').bold = True

    if user_current.ruSubject9:
        paragraph3.add_run('\n- ' + user_current.ruSubject9)
        paragraph3.add_run(' (' + user_current.viSubject9 + '): ')
        paragraph3.add_run(user_current.resultSubject9 + ' (' + GRADE_CHOICES[user_current.resultSubject9] + ')').bold = True

    if user_current.ruSubject10:
        paragraph3.add_run('\n- ' + user_current.ruSubject10)
        paragraph3.add_run(' (' + user_current.viSubject10 + '): ')
        paragraph3.add_run(user_current.resultSubject10 + ' (' + GRADE_CHOICES[user_current.resultSubject10] + ')').bold = True

    if user_current.ruSubject11:
        paragraph3.add_run('\n- ' + user_current.ruSubject11)
        paragraph3.add_run(' (' + user_current.viSubject11 + '): ')
        paragraph3.add_run(user_current.resultSubject11 + ' (' + GRADE_CHOICES[user_current.resultSubject11] + ')').bold = True

    if user_current.ruSubject12:
        paragraph3.add_run('\n- ' + user_current.ruSubject12)
        paragraph3.add_run(' (' + user_current.viSubject12 + '): ')
        paragraph3.add_run(user_current.resultSubject12 + ' (' + GRADE_CHOICES[user_current.resultSubject12] + ')').bold = True

    paragraph3.add_run('\n\nMột đơn vị học trình gồm ... giờ.')

    table2 = document.add_table(rows=1, cols=2)
    row = table2.rows[0]

    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].paragraphs[0].add_run('Trưởng khoa\n\nChernhikov A.C.\n(Đã kí và đóng dấu)')

    table3 = document.add_table(rows=1, cols=2)
    row = table3.rows[0]

    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[1].paragraphs[0].add_run('Tôi xin cam đoan dịch đúng so với bản gốc.').bold = True
    row.cells[1].paragraphs[0].add_run('\nMát-xcơ-va, ngày  tháng  năm 20  ')
    row.cells[1].paragraphs[0].add_run('\nNGƯỜI DỊCH')
    row.cells[1].paragraphs[0].add_run('\n(ký và ghi rõ họ tên)')
    row.cells[1].paragraphs[0].add_run('\n\n\n\n' + user_current.fullName).bold = True

    # Prepare document for download        
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length

    return response